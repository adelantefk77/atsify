from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT_NAME = "Calibri"


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.bold = True
    # Bottom border under the heading, matching the PDF renderer's section rule.
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _para(doc, text, bold=False, italic=False, size=11, space_after=2, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text or "")
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def _bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text or "")
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    return p


def render_cv_docx(cv_data: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    contact = cv_data.get("contact") or {}

    _heading(doc, "Dane kontaktowe")

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run(contact.get("name") or "")
    run_name.font.name = FONT_NAME
    run_name.font.size = Pt(20)
    run_name.bold = True

    if contact.get("professional_title"):
        _para(doc, contact["professional_title"], size=13, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    contact_line = " | ".join(filter(None, [contact.get("phone"), contact.get("email"), contact.get("location")]))
    if contact_line:
        _para(doc, contact_line, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if contact.get("linkedin"):
        _para(doc, contact["linkedin"], size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _heading(doc, "Podsumowanie zawodowe")
    _para(doc, cv_data.get("summary"))

    experience = cv_data.get("experience") or []
    if experience:
        _heading(doc, "Doświadczenie zawodowe")
        for job in experience:
            meta = " | ".join(filter(None, [job.get("company"), job.get("location"), job.get("dates")]))
            _para(doc, job.get("title"), bold=True, space_after=0)
            if meta:
                _para(doc, meta, italic=True, size=10, space_after=3)
            for bullet in job.get("bullets") or []:
                _bullet(doc, bullet)

    education = cv_data.get("education") or []
    if education:
        _heading(doc, "Wykształcenie")
        for edu in education:
            meta = " | ".join(filter(None, [edu.get("school"), edu.get("location"), edu.get("dates")]))
            _para(doc, edu.get("degree"), bold=True, space_after=0)
            if meta:
                _para(doc, meta, italic=True, size=10, space_after=3)
            for detail in edu.get("details") or []:
                _bullet(doc, detail)

    skills = cv_data.get("skills") or []
    if skills:
        _heading(doc, "Umiejętności")
        for group in skills:
            items = ", ".join(group.get("items") or [])
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run_cat = p.add_run(f"{group.get('category') or ''}: ")
            run_cat.font.name = FONT_NAME
            run_cat.font.size = Pt(11)
            run_cat.bold = True
            run_items = p.add_run(items)
            run_items.font.name = FONT_NAME
            run_items.font.size = Pt(11)

    certifications = cv_data.get("certifications") or []
    if certifications:
        _heading(doc, "Certyfikaty / szkolenia")
        for c in certifications:
            _bullet(doc, c)

    languages = cv_data.get("languages") or []
    if languages:
        _heading(doc, "Języki obce")
        for lang in languages:
            _bullet(doc, lang)

    _heading(doc, "Klauzula RODO")
    _para(doc, cv_data.get("rodo_clause"), size=8, italic=True)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
